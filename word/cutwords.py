import os
import jieba
import docx
file = docx.Document('123.docx')
# with open('result.txt', mode='w', encoding='utf-8') as f:
#     file = docx.Document('123.docx')
#     for para in file.paragraphs:
#         f.write(para.text)
# seg_list = jieba.cut(text, cut_all=False)
# print('reslult:'+'/'.join(seg_list))
